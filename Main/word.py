# word.py
"""
Word document generation from map data using docxtpl templates.
"""

import os
import io

from .config import cfg, CACHE_DIR
from .utils import number_to_bulan, dasarian_romawi
from .static import redownload
from .status import update as status_update
from .narasi import get_full_narration, build_table_data


def _build_table_subdoc(doc, table_data):
    """Create a docxtpl subdocument containing a formatted table.

    Args:
        doc: DocxTemplate instance (needed for new_subdoc).
        table_data: dict with 'columns' and 'rows' from build_table_data().

    Returns:
        Subdoc object to embed in the template context.
    """
    from docx.shared import Pt
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    sd = doc.new_subdoc()

    num_cols = len(table_data['columns'])
    num_rows = len(table_data['rows']) + 1  # +1 for header

    table = sd.add_table(rows=num_rows, cols=num_cols)

    # Apply borders via XML (subdocuments don't have built-in Word styles)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)

    # ---- Header row ----
    for j, col_name in enumerate(table_data['columns']):
        cell = table.cell(0, j)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(col_name)
        run.bold = True
        run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Shade header cells light blue
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    # ---- Data rows ----
    for i, row_data in enumerate(table_data['rows']):
        for j, val in enumerate(row_data):
            cell = table.cell(i + 1, j)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)

    return sd


def _print_fallback(title1, title2, analysis, table_data):
    """Print narration and table to console as a fallback."""
    has_content = analysis or table_data
    if not has_content:
        return

    print("\n" + "=" * 60)
    print(f"  {title1} {title2}")
    print("=" * 60)

    if analysis:
        print(f"\n{analysis}\n")

    if table_data:
        columns = table_data['columns']
        rows = table_data['rows']
        # Compute column widths
        col_widths = [len(c) for c in columns]
        for row in rows:
            for j, val in enumerate(row):
                col_widths[j] = max(col_widths[j], len(val))

        def fmt_row(vals):
            return ' | '.join(v.ljust(col_widths[j]) for j, v in enumerate(vals))

        header = fmt_row(columns)
        print(header)
        print('-+-'.join('-' * w for w in col_widths))
        for row in rows:
            print(fmt_row(row))
        print()


def arrange_word(map_data):
    """Generate a Word document report from map data.

    Args:
        map_data: dict returned by execute(), must contain 'image' (PIL Image)
                  and map metadata (peta, tipe, skala, year, month, etc.).

    Returns:
        str: Output file path on success, None on failure.
    """
    peta = map_data['peta']
    tipe = map_data['tipe']
    skala = map_data['skala']
    year = map_data['year']
    month = map_data['month']

    # Build title and period strings
    title1 = f'{peta} {tipe}'
    if skala == 'Bulanan':
        title2 = f'Bulan {number_to_bulan(month)} {year}'
    else:
        title2 = (f'Bulan {number_to_bulan(month)} '
                   f'Dasarian {dasarian_romawi(map_data["dasarian"])} {year}')

    # Compute narration and table data before any docx work so they
    # are available for the console fallback if Word generation fails.
    analysis = None
    table_data = None
    try:
        analysis = get_full_narration(map_data)
        table_data = build_table_data(map_data)
    except Exception:
        pass  # best-effort; failures handled in fallback below

    try:
        try:
            from docxtpl import DocxTemplate, InlineImage
            import docxcompose  # noqa: F401 – verify dependency is available
        except ImportError:
            import subprocess
            status_update("Installing docxtpl...")
            subprocess.check_call(
                ['pip', 'install', 'docxtpl', 'docxcompose', '-q']
            )
            from docxtpl import DocxTemplate, InlineImage
        from docx.shared import Cm
        from google.colab import files

        status_update("Generating Word document...")

        desc = f'Peta {title1} {title2}'

        # Convert PIL Image to BytesIO buffer
        image_buffer = io.BytesIO()
        map_data['image'].save(image_buffer, format='PNG')
        image_buffer.seek(0)

        # Load template and build context
        template_path = os.path.join(CACHE_DIR, 'template_doc.docx')
        try:
            doc = DocxTemplate(template_path)
        except Exception:
            status_update("template_doc.docx is missing or corrupted, re-downloading")
            redownload("template_doc.docx")
            doc = DocxTemplate(template_path)

        # Build table for text2 field (if applicable)
        if table_data:
            text2_content = _build_table_subdoc(doc, table_data)
        else:
            text2_content = ''

        context = {
            'title1': f'{title1} {title2}',
            'image1': InlineImage(doc, image_buffer, width=Cm(15)),
            'desc': desc,
            'text1': analysis or '',
            'text2': text2_content,
        }

        # Render and save
        doc.render(context)
        output_path = f'/content/Laporan {peta} {tipe}.docx'
        doc.save(output_path)

        status_update(f"Word exported: {os.path.basename(output_path)}")
        files.download(output_path)

        return output_path

    except Exception as e:
        try:
            from google.genai.errors import ServerError
            if isinstance(e, ServerError):
                print(
                    "Error generating Word document: Gemini API unavailable "
                    "after retries. Try again in a few minutes."
                )
                _print_fallback(title1, title2, analysis, table_data)
                return None
        except ImportError:
            pass
        print(f"Error generating Word document: {e}")
        _print_fallback(title1, title2, analysis, table_data)
        return None
